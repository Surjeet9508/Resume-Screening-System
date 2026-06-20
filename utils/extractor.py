from io import BytesIO
import re
import PyPDF2
import docx


def clean_extracted_text(text: str) -> str:
    """
    Clean extracted text lightly without removing useful resume content.
    """
    if not text:
        return ""

    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n", text)
    return text.strip()


def extract_pdf(file):
    """
    Extract text from text-based PDF files safely.
    Returns empty string if extraction fails.
    """
    try:
        file.seek(0)
        pdf_bytes = file.read()
        file.seek(0)

        reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
        text_parts = []

        for page in reader.pages:
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
            except Exception:
                continue

        final_text = "\n".join(text_parts)
        return clean_extracted_text(final_text)

    except Exception:
        return ""


def extract_docx(file):
    """
    Extract text from DOCX safely.
    Includes paragraphs, tables, headers, and footers.
    Returns empty string if extraction fails.
    """
    try:
        file.seek(0)
        document = docx.Document(BytesIO(file.read()))
        file.seek(0)

        text_parts = []

        # Paragraphs
        for para in document.paragraphs:
            para_text = para.text.strip()
            if para_text:
                text_parts.append(para_text)

        # Tables
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))

        # Headers and footers
        for section in document.sections:
            try:
                for para in section.header.paragraphs:
                    para_text = para.text.strip()
                    if para_text:
                        text_parts.append(para_text)
            except Exception:
                pass

            try:
                for para in section.footer.paragraphs:
                    para_text = para.text.strip()
                    if para_text:
                        text_parts.append(para_text)
            except Exception:
                pass

        final_text = "\n".join(text_parts)
        return clean_extracted_text(final_text)

    except Exception:
        return ""